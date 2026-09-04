"""Private document library and lightweight local hybrid retrieval."""
from __future__ import annotations

import hashlib, io, math, os, re, secrets, shutil, threading, time, warnings, zipfile
from collections import Counter, defaultdict
from pathlib import Path
from typing import Literal

from pydantic import BaseModel, ConfigDict, Field

from . import store

MAX_BYTES = 50 * 1024 * 1024
MAX_TEXT = 2_500_000
MAX_CHUNKS = 2500
MODEL_NAME = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
MODEL_REPO = "qdrant/paraphrase-multilingual-MiniLM-L12-v2-onnx-Q"
MODEL_REVISION = "faf4aa4225822f3bc6376869cb1164e8e3feedd0"
MODEL_FILES = ("config.json", "model_optimized.onnx", "special_tokens_map.json", "tokenizer_config.json", "tokenizer.json")
MODEL_CACHE = "models/rag"
_MODEL = None
_MODEL_LOCK = threading.RLock()


class DocumentEdit(BaseModel):
    model_config = ConfigDict(extra="forbid", str_strip_whitespace=True)
    title: str = Field(min_length=1, max_length=180)
    author: str = Field("", max_length=160)
    year: str = Field("", max_length=40)
    provenance: str = Field("", max_length=500)
    enabled: bool = True


class PastedDocument(DocumentEdit):
    text: str = Field(min_length=80, max_length=MAX_TEXT)


def folder(did: str) -> Path:
    if not re.fullmatch(r"[a-f0-9]{24}", did):
        raise KeyError(did)
    return store.DATA / "documents" / did


def get(did: str):
    path = folder(did) / "record.json"
    if not path.is_file():
        raise KeyError(did)
    return store.read_json(path)


def catalog():
    with store.LOCK:
        root = store.DATA / "documents"
        return sorted((store.read_json(p) for p in root.glob("*/record.json")), key=lambda r: r["created"], reverse=True)


def validate_selection(ids: list[str], enabled: bool = True):
    records = [get(did) for did in dict.fromkeys(ids)]
    if enabled:
        unavailable = [r["title"] for r in records if not r.get("enabled", True)]
        scans = [r["title"] for r in records if r.get("status") == "needs_ocr"]
        if unavailable:
            raise ValueError("Documento non disponibile: " + ", ".join(unavailable[:3]) + ".")
        if scans:
            raise ValueError("Serve una versione con OCR per usare: " + ", ".join(scans[:3]) + ".")
    return records


def save(did: str, value: DocumentEdit):
    with store.LOCK:
        record = get(did)
        record.update(value.model_dump())
        record["updated"] = store.now()
        store.write_json(folder(did) / "record.json", record)
        return record


def _safe_name(filename: str):
    name = filename.replace("\\", "/").rsplit("/", 1)[-1][:180]
    return "".join(c for c in name if ord(c) >= 32) or "documento.txt"


def _safe_docx(raw: bytes):
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            entries = archive.infolist()
            if len(entries) > 5000 or sum(x.file_size for x in entries) > 120 * 1024 * 1024:
                raise ValueError("Il documento Word è troppo grande una volta aperto.")
            if not any(x.filename == "word/document.xml" for x in entries):
                raise ValueError("File DOCX non valido.")
    except zipfile.BadZipFile as error:
        raise ValueError("File DOCX non valido.") from error


def extract_file(raw: bytes, filename: str):
    name = _safe_name(filename)
    suffix = Path(name).suffix.lower()
    if suffix in (".txt", ".md"):
        try:
            text = raw.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise ValueError("Il file di testo deve essere UTF-8.") from error
        pages = [{"page": 1, "text": text}]
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
            # Many valid historical scans contain minor structural errors.  The
            # permissive parser still reads them locally and never executes PDF
            # actions or embedded files.
            reader = PdfReader(io.BytesIO(raw), strict=False)
            if reader.is_encrypted:
                raise ValueError("Il PDF è protetto da password.")
            if len(reader.pages) > 2000:
                raise ValueError("Il PDF supera 2000 pagine.")
            pages = []
            extracted = 0
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                extracted += len(text)
                if extracted > MAX_TEXT:
                    raise ValueError("Il documento supera il limite di testo indicizzabile.")
                pages.append({"page": i + 1, "text": text})
        except ValueError:
            raise
        except Exception as error:
            raise ValueError("PDF non valido o non leggibile.") from error
    elif suffix == ".docx":
        _safe_docx(raw)
        try:
            from docx import Document
            document = Document(io.BytesIO(raw))
            lines = [p.text for p in document.paragraphs if p.text.strip()]
            # Dates, names and figures in source documents are often stored in
            # tables; include them in retrieval instead of silently dropping
            # them.
            for table in document.tables:
                for row in table.rows:
                    value = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if value:
                        lines.append(value)
            text = "\n".join(lines)
            pages = [{"page": None, "section": "Documento Word", "text": text}]
        except Exception as error:
            raise ValueError("Documento Word non valido o non leggibile.") from error
    elif suffix == ".doc":
        raise ValueError("Il formato DOC precedente al 2007 non è leggibile in modo sicuro: salvalo come DOCX.")
    else:
        raise ValueError("Formati supportati: PDF, DOCX, TXT e Markdown.")
    total = sum(len(p["text"]) for p in pages)
    if total > MAX_TEXT:
        raise ValueError("Il documento supera il limite di testo indicizzabile.")
    return name, pages


def chunks(pages):
    result = []
    for page in pages:
        clean = re.sub(r"[ \t]+", " ", page.get("text", "")).replace("\r", "").strip()
        blocks = [x.strip() for x in re.split(r"\n{2,}|(?<=[.!?])\s+(?=[A-ZÀ-ÖØ-Ý])", clean) if x.strip()]
        current = ""
        for block in blocks:
            if len(current) + len(block) + 1 <= 1350:
                current = (current + " " + block).strip()
                continue
            if current:
                result.append({"page": page.get("page"), "section": page.get("section", ""), "text": current})
            current = (current[-180:] + " " + block).strip() if current else block
            while len(current) > 1600:
                result.append({"page": page.get("page"), "section": page.get("section", ""), "text": current[:1350]})
                current = current[1170:]
        if current:
            result.append({"page": page.get("page"), "section": page.get("section", ""), "text": current})
        if len(result) > MAX_CHUNKS:
            raise ValueError("Il documento produce troppi passaggi da indicizzare.")
    return [x for x in result if len(x["text"]) >= 40]


def embedding_model(download: bool = False):
    global _MODEL
    with _MODEL_LOCK:
        if _MODEL is not None:
            return _MODEL
        cache = store.DATA / MODEL_CACHE
        if not download and not (cache / "ready.json").is_file():
            raise RuntimeError("Modello RAG non installato: riapri INSTALLA.bat.")
        cache.mkdir(parents=True, exist_ok=True)
        os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")
        snapshot = cache / ("models--" + MODEL_REPO.replace("/", "--")) / "snapshots" / MODEL_REVISION
        if download and not all((snapshot / name).is_file() for name in MODEL_FILES):
            from huggingface_hub import snapshot_download
            error = None
            for attempt in range(3):
                try:
                    snapshot = Path(snapshot_download(repo_id=MODEL_REPO, revision=MODEL_REVISION,
                                                       allow_patterns=list(MODEL_FILES), cache_dir=str(cache)))
                    error = None
                    break
                except (OSError, RuntimeError, ValueError) as caught:
                    error = caught
                    if attempt < 2:
                        time.sleep(2 ** attempt)
            if error is not None:
                raise RuntimeError("Download del modello RAG non completato; riapri INSTALLA.bat.") from error
        if not all((snapshot / name).is_file() for name in MODEL_FILES):
            raise RuntimeError("File del modello RAG incompleti: riapri INSTALLA.bat.")
        from fastembed import TextEmbedding
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", message=r"The model .* now uses mean pooling.*")
            _MODEL = TextEmbedding(model_name=MODEL_NAME, cache_dir=str(cache), specific_model_path=str(snapshot),
                                   threads=max(1, min(4, os.cpu_count() or 1)))
        return _MODEL


def ensure_model(download: bool = False):
    model = embedding_model(download)
    vector = list(model.embed(["Una fonte storica locale."]))[0]
    if len(vector) != 384:
        raise RuntimeError("Il modello RAG ha restituito vettori non validi.")
    if download:
        files = []
        root = store.DATA / MODEL_CACHE
        for path in sorted(root.rglob("*")):
            if path.is_file() and path.name != "ready.json":
                files.append({"path": path.relative_to(root).as_posix(), "bytes": path.stat().st_size,
                              "sha256": hashlib.sha256(path.read_bytes()).hexdigest()})
        store.write_json(root / "ready.json", {"model": MODEL_NAME, "repository": MODEL_REPO,
                                               "revision": MODEL_REVISION, "dimensions": 384, "files": files})
    return True


def index_document(did: str):
    record = get(did)
    if record.get("status") == "needs_ocr":
        return record
    rows = store.read_json(folder(did) / "chunks.json")
    try:
        import numpy as np
        model = embedding_model(False)
        values = np.asarray(list(model.embed([x["text"] for x in rows])), dtype=np.float32)
        if values.shape != (len(rows), 384) or not np.isfinite(values).all():
            raise RuntimeError("Indice semantico non valido.")
        norms = np.linalg.norm(values, axis=1, keepdims=True)
        values = values / np.maximum(norms, 1e-12)
        target = folder(did) / "vectors.npy"
        temporary = target.with_suffix(".npy.part")
        with temporary.open("wb") as output:
            np.save(output, values.astype(np.float16), allow_pickle=False)
        temporary.replace(target)
        record.update(status="indexed", index_model=MODEL_NAME, index_error="")
    except Exception as error:
        record.update(status="text_ready", index_model="lexical", index_error=str(error)[:300])
    record["updated"] = store.now()
    store.write_json(folder(did) / "record.json", record)
    return record


def _store(raw: bytes, filename: str, pages, value: DocumentEdit):
    did = secrets.token_hex(12)
    dest = folder(did)
    dest.mkdir(parents=True)
    name = _safe_name(filename)
    original = "original" + (Path(name).suffix.lower() or ".txt")
    (dest / original).write_bytes(raw)
    rows = chunks(pages)
    chars = sum(len(x["text"]) for x in rows)
    status = "needs_ocr" if Path(name).suffix.lower() == ".pdf" and chars < 400 else "text_ready"
    record = value.model_dump()
    record.update(id=did, filename=name, original=original, created=store.now(), updated=store.now(),
                  bytes=len(raw), characters=chars, chunks=len(rows), pages=len(pages), status=status,
                  sha256=hashlib.sha256(raw).hexdigest(), index_model="", index_error="")
    store.write_json(dest / "chunks.json", rows)
    store.write_json(dest / "record.json", record)
    return index_document(did) if status != "needs_ocr" else record


def upload(raw: bytes, filename: str):
    if not raw or len(raw) > MAX_BYTES:
        raise ValueError("Usa un documento fino a 50 MB.")
    name, pages = extract_file(raw, filename)
    return _store(raw, name, pages, DocumentEdit(title=Path(name).stem or "Documento"))


def paste(value: PastedDocument):
    raw = value.text.encode("utf-8")
    meta = DocumentEdit(**value.model_dump(exclude={"text"}))
    return _store(raw, value.title + ".txt", [{"page": 1, "section": "Testo inserito", "text": value.text}], meta)


def freeze(pid: str, ids: list[str], enabled: bool = True):
    checkpoint = store.JOBS / pid / "checkpoints" / "document-selection.json"
    with store.LOCK:
        if checkpoint.exists():
            return store.read_json(checkpoint)["documents"]
        selected = []
        if enabled:
            for did in dict.fromkeys(ids):
                item = get(did)
                if not item.get("enabled", True) or item.get("status") == "needs_ocr":
                    continue
                source = folder(did)
                destination = store.JOBS / pid / "workspace" / "assets" / "documents" / did
                destination.mkdir(parents=True, exist_ok=True)
                if hashlib.sha256((source / item["original"]).read_bytes()).hexdigest() != item["sha256"]:
                    raise ValueError("Un documento è stato modificato sul disco: caricalo nuovamente.")
                for name in (item["original"], "chunks.json", "record.json", "vectors.npy"):
                    if (source / name).is_file():
                        shutil.copy2(source / name, destination / name)
                selected.append(item)
        store.write_json(checkpoint, {"created": store.now(), "documents": selected})
        return selected


def _tokens(text: str):
    return re.findall(r"[a-zà-öø-ÿ0-9]{2,}", text.casefold())


def retrieve(pid: str, query: str, limit: int = 12):
    snapshot = store.JOBS / pid / "checkpoints" / "document-selection.json"
    if not snapshot.is_file():
        return []
    docs = store.read_json(snapshot)["documents"]
    candidates = []
    for doc in docs:
        root = store.JOBS / pid / "workspace" / "assets" / "documents" / doc["id"]
        for index, row in enumerate(store.read_json(root / "chunks.json")):
            candidates.append({"doc": doc, "root": root, "index": index, **row})
    if not candidates:
        return []
    terms = _tokens(query)
    corpus = [Counter(_tokens(x["text"])) for x in candidates]
    average = sum(sum(x.values()) for x in corpus) / max(1, len(corpus))
    document_frequency = Counter(term for counts in corpus for term in counts)
    lexical = []
    for counts in corpus:
        length = sum(counts.values())
        score = 0.0
        for term in terms:
            tf = counts[term]
            if not tf:
                continue
            idf = math.log(1 + (len(corpus) - document_frequency[term] + .5) / (document_frequency[term] + .5))
            score += idf * tf * 2.2 / (tf + 1.2 * (.25 + .75 * length / max(1, average)))
        lexical.append(score)
    dense = [0.0] * len(candidates)
    try:
        import numpy as np
        query_vector = np.asarray(list(embedding_model(False).embed([query]))[0], dtype=np.float32)
        query_vector /= max(float(np.linalg.norm(query_vector)), 1e-12)
        offset = 0
        for doc in docs:
            count = sum(1 for x in candidates if x["doc"]["id"] == doc["id"])
            path = store.JOBS / pid / "workspace" / "assets" / "documents" / doc["id"] / "vectors.npy"
            if path.is_file():
                vectors = np.load(path, allow_pickle=False).astype(np.float32)
                if len(vectors) == count:
                    dense[offset:offset + count] = ((vectors @ query_vector) + 1).tolist()
            offset += count
    except Exception:
        pass
    lexical_max = max(lexical) or 1.0
    for i, item in enumerate(candidates):
        item["score"] = .45 * lexical[i] / lexical_max + .55 * dense[i] / 2
    chosen = []
    per_document = Counter()
    for item in sorted(candidates, key=lambda x: x["score"], reverse=True):
        if per_document[item["doc"]["id"]] >= 4:
            continue
        chosen.append(item)
        per_document[item["doc"]["id"]] += 1
        if len(chosen) >= limit:
            break
    grouped = defaultdict(list)
    for item in chosen:
        grouped[item["doc"]["id"]].append(item)
    sources = []
    ranked_documents = sorted(enumerate(docs, 1),
                              key=lambda pair: max((r["score"] for r in grouped.get(pair[1]["id"], [])), default=-1),
                              reverse=True)
    for number, doc in ranked_documents:
        rows = grouped.get(doc["id"])
        if not rows:
            continue
        extracts = []
        for row in sorted(rows, key=lambda x: x["index"]):
            where = f'pagina {row["page"]}' if row.get("page") else row.get("section") or f'passaggio {row["index"] + 1}'
            extracts.append(f"[{where}]\n{row['text']}")
        author = doc.get("author") or "autore non indicato"
        year = (", " + doc["year"]) if doc.get("year") else ""
        citation = "; ".join(x for x in (f"{author}{year}", doc.get("provenance", ""), doc["filename"]) if x)
        sources.append({"id": f"D{number}", "title": doc["title"],
                        "url": f"assets/documents/{doc['id']}/{doc['original']}",
                        "text": "\n\n".join(extracts), "retrieved": store.now(), "sha256": doc["sha256"],
                        "origin": "local_document", "citation": citation,
                        "document_id": doc["id"]})
    return sources
