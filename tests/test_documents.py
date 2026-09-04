import io, json, zipfile

import numpy as np
import pytest
from fastapi.testclient import TestClient

from app import documents, server, store
from app.models import ProjectRequest


class FakeEmbeddings:
    def embed(self, texts):
        for text in texts:
            value = np.zeros(384, dtype=np.float32)
            for word in documents._tokens(text):
                value[hash(word) % 384] += 1
            norm = np.linalg.norm(value)
            yield value / max(float(norm), 1.0)


@pytest.fixture(autouse=True)
def isolated(tmp_path, monkeypatch):
    jobs = tmp_path / "jobs"
    jobs.mkdir()
    for module in (store, server):
        monkeypatch.setattr(module, "DATA", tmp_path)
        monkeypatch.setattr(module, "JOBS", jobs)
    monkeypatch.setattr(documents, "embedding_model", lambda download=False: FakeEmbeddings())
    store.init()


@pytest.fixture
def client():
    return TestClient(server.app, headers={"X-DocumentariAI": "studio"})


def long_text(subject="Waterloo"):
    return ((subject + " è discusso nella fonte con cronologia, luoghi e protagonisti verificabili. ") * 35).strip()


def test_text_upload_index_snapshot_and_hybrid_retrieval(tmp_path):
    waterloo = documents.upload(long_text().encode(), "waterloo.txt")
    rome = documents.paste(documents.PastedDocument(title="Roma", author="Storico", year="2024", text=long_text("Roma antica")))
    assert waterloo["status"] == rome["status"] == "indexed"
    assert (documents.folder(waterloo["id"]) / "vectors.npy").is_file()
    project = store.create(ProjectRequest(topic="Battaglia di Waterloo", start=False, document_ids=[rome["id"], waterloo["id"]]))
    documents.freeze(project["id"], project["document_ids"])
    sources = documents.retrieve(project["id"], "Waterloo battaglia", limit=1)
    assert sources[0]["origin"] == "local_document" and sources[0]["id"].startswith("D")
    assert len(sources) == 1 and "Waterloo" in sources[0]["text"] and sources[0]["url"].startswith("assets/documents/")
    snapshot = store.JOBS / project["id"] / "workspace" / sources[0]["url"]
    assert snapshot.is_file() and snapshot.read_bytes() == long_text().encode()


def test_document_api_paste_metadata_and_project_selection(client):
    response = client.post("/api/documents/text", json={"title": "Fonte locale", "author": "A. Test", "year": "1900", "provenance": "Archivio", "enabled": True, "text": long_text()})
    assert response.status_code == 201, response.text
    doc = response.json()
    assert client.get("/api/documents").json()[0]["id"] == doc["id"]
    assert client.get(f"/api/documents/{doc['id']}/file").status_code == 200
    project = client.post("/api/projects", json={"topic": "Battaglia di Waterloo", "start": False, "document_ids": []}).json()
    result = client.put(f"/api/projects/{project['id']}/documents", json={"enabled": True, "document_ids": [doc["id"]]})
    assert result.status_code == 200 and result.json()["selected_ids"] == [doc["id"]]
    assert client.put(f"/api/projects/{project['id']}/documents", json={"enabled": True, "document_ids": ["../settings"]}).status_code in (404, 422)
    documents.freeze(project["id"], [doc["id"]])
    source = documents.retrieve(project["id"], "Waterloo", limit=1)[0]
    assert "A. Test, 1900; Archivio;" in source["citation"]
    exported = client.get(f"/api/projects/{project['id']}/export")
    assert exported.status_code == 200
    with zipfile.ZipFile(io.BytesIO(exported.content)) as archive:
        assert any(name.endswith("/original.txt") for name in archive.namelist())


def test_pdf_without_text_is_kept_with_clear_ocr_status(client):
    from pypdf import PdfWriter
    stream = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=200)
    writer.write(stream)
    record = documents.upload(stream.getvalue(), "scansione.pdf")
    assert record["status"] == "needs_ocr" and record["pages"] == 1
    response = client.post("/api/projects", json={"topic": "Tema dalla scansione", "start": False,
                                                   "document_ids": [record["id"]]})
    assert response.status_code == 400 and "OCR" in response.json()["detail"]
    assert client.get("/api/projects").json() == []


def test_docx_text_is_extracted():
    from docx import Document
    stream = io.BytesIO()
    value = Document()
    value.add_heading("Campagna napoleonica", 1)
    value.add_paragraph(long_text("Austerlitz"))
    table = value.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Napoleone"
    table.cell(0, 1).text = "2 dicembre 1805"
    value.save(stream)
    record = documents.upload(stream.getvalue(), "fonte.docx")
    assert record["status"] == "indexed" and record["characters"] > 1000
    assert "2 dicembre 1805" in json.dumps(store.read_json(documents.folder(record["id"]) / "chunks.json"))


def test_legacy_doc_and_untrusted_requests_are_rejected(client):
    assert client.post("/api/documents/file?filename=vecchio.doc", content=b"legacy").status_code == 400
    assert TestClient(server.app).post("/api/documents/text", json={"title": "x", "text": long_text()}).status_code == 403
    assert client.post("/api/projects", json={"topic": "Tema valido", "start": False, "document_ids": ["not-safe"]}).status_code == 422
